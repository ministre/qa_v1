from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from django.shortcuts import get_object_or_404
from protocol.models import Protocol
from testplan.models import TestPlan, Category, Test, TestImage, TestConfig
from django.utils.translation import gettext_lazy as _
from docx import Document
from django.conf import settings
import os
from django.http import HttpResponse, Http404
from django.utils.datastructures import MultiValueDictKeyError
from docx.shared import Cm, Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from django.views.generic import CreateView, ListView, UpdateView, DeleteView
from django.utils.decorators import method_decorator
from .models import DocxProfile
from .forms import DocxProfileForm
from django.urls import reverse
from django.utils import timezone
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn


@method_decorator(login_required, name='dispatch')
class DocxProfileListView(ListView):
    context_object_name = 'docx_profiles'
    queryset = DocxProfile.objects.all().order_by('id')
    template_name = 'docx_builder/docx_profiles.html'


@method_decorator(login_required, name='dispatch')
class DocxProfileCreate(CreateView):
    model = DocxProfile
    form_class = DocxProfileForm
    template_name = 'device/create.html'

    def get_initial(self):
        return {'created_by': self.request.user, 'updated_by': self.request.user}

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['back_url'] = reverse('docx_profiles')
        return context

    def get_success_url(self):
        return reverse('docx_profiles')


@method_decorator(login_required, name='dispatch')
class DocxProfileUpdate(UpdateView):
    model = DocxProfile
    form_class = DocxProfileForm
    template_name = 'docx_builder/update.html'

    def get_initial(self):
        return {'updated_by': self.request.user, 'updated_at': timezone.now()}

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['back_url'] = reverse('docx_profiles')
        return context

    def get_success_url(self):
        return reverse('docx_profiles')


@method_decorator(login_required, name='dispatch')
class DocxProfileDelete(DeleteView):
    model = DocxProfile
    template_name = 'device/delete.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['back_url'] = reverse('docx_profiles')
        return context

    def get_success_url(self):
        return reverse('docx_profiles')


@login_required
def build_protocol(request):
    if request.method == 'POST':
        protocol = get_object_or_404(Protocol, id=request.POST['protocol_id'])
        docx_profile = get_object_or_404(DocxProfile, id=request.POST['docx_profile_id'])
        header = general = performance = results_table = summary = team = False
        try:
            if request.POST['header']:
                header = True
        except MultiValueDictKeyError:
            pass
        try:
            if request.POST['general']:
                general = True
        except MultiValueDictKeyError:
            pass
        try:
            if request.POST['performance']:
                performance = True
        except MultiValueDictKeyError:
            pass
        try:
            if request.POST['results_table']:
                results_table = True
        except MultiValueDictKeyError:
            pass
        try:
            if request.POST['summary']:
                summary = True
        except MultiValueDictKeyError:
            pass
        try:
            if request.POST['team']:
                team = True
        except MultiValueDictKeyError:
            pass

        document = build_document(docx_profile=docx_profile)

        # margin
        section = document.sections
        section[0].left_margin = Cm(2)
        section[0].right_margin = Cm(1)
        section[0].top_margin = Cm(1)
        section[0].bottom_margin = Cm(1)
        ###
        if header:
            style = document.styles.add_style('Header Table', WD_STYLE_TYPE.TABLE)
            style.base_style = document.styles['Table Grid']

            header = document.sections[0].header
            table = header.add_table(rows=0, cols=3, width=Cm(19.5))
            table.style = 'Header Table'
            row_cells = table.add_row().cells
            paragraph = row_cells[0].paragraphs[0]
            run = paragraph.add_run()
            if docx_profile.header_logo:
                run.add_picture(docx_profile.header_logo, width=Inches(1.25))
            paragraph = row_cells[1].paragraphs[0]
            run = paragraph.add_run()
            run.add_text(str(_('Protocol')) + ' ' + protocol.device.vendor.name + ' ' + protocol.device.model)
            if protocol.device.hw:
                run.add_text(' (' + protocol.device.hw + ')')
            row_cells = table.add_row().cells
            paragraph = row_cells[0].paragraphs[0]
            run = paragraph.add_run()
            if docx_profile.header_text1:
                run.add_text(docx_profile.header_text1)
            paragraph = row_cells[1].paragraphs[0]
            run = paragraph.add_run()
            if docx_profile.header_text2:
                run.add_text(docx_profile.header_text2)

            table.cell(0, 1).merge(table.cell(0, 2))
            table.cell(0, 0).width = Cm(5)
            table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.cell(1, 0).width = Cm(5)
            table.cell(1, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.cell(0, 1).width = Cm(10.5)
            table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.cell(1, 1).width = Cm(10.5)
            table.cell(1, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.cell(0, 2).width = Cm(4)
            table.cell(0, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.cell(1, 2).width = Cm(4)
            table.cell(1, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.cell(0, 1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        if general:
            document.add_paragraph(str(_('General Information about DUT')), style='Heading 1')
            table = document.add_table(rows=0, cols=2)
            table.style = 'TableGrid'
            row_cells = table.add_row().cells
            row_cells[0].text = str(_('Device Type')) + ': '
            row_cells[1].text = protocol.device.type.name
            row_cells = table.add_row().cells
            row_cells[0].text = str(_('Vendor')) + ': '
            row_cells[1].text = protocol.device.vendor.name
            row_cells = table.add_row().cells
            row_cells[0].text = str(_('Model')) + ': '
            row_cells[1].text = protocol.device.model
            if protocol.device.hw:
                row_cells = table.add_row().cells
                row_cells[0].text = str(_('Hardware Version')) + ': '
                row_cells[1].text = protocol.device.hw
            row_cells = table.add_row().cells
            row_cells[0].text = str(_('Software Version')) + ': '
            row_cells[1].text = protocol.sw
            if protocol.sw_checksum:
                row_cells = table.add_row().cells
                row_cells[0].text = str(_('Software Checksum')) + ': '
                row_cells[1].text = protocol.sw_checksum
            row_cells = table.add_row().cells
            row_cells[0].merge(row_cells[1])
            row_cells[0].text = str(_('Hardware Specifications')) + ': '
            row_cells[0].paragraphs[0].runs[0].font.bold = True
            row_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            if protocol.device.interfaces:
                row_cells = table.add_row().cells
                row_cells[0].text = str(_('Interfaces')) + ': '
                row_cells[1].text = protocol.device.interfaces
            if protocol.device.leds:
                row_cells = table.add_row().cells
                row_cells[0].text = str(_('Leds')) + ': '
                row_cells[1].text = protocol.device.leds
            if protocol.device.buttons:
                row_cells = table.add_row().cells
                row_cells[0].text = str(_('Buttons')) + ': '
                row_cells[1].text = protocol.device.buttons
            if protocol.device.chipsets:
                row_cells = table.add_row().cells
                row_cells[0].text = str(_('Chipsets')) + ': '
                row_cells[1].text = protocol.device.chipsets
            if protocol.device.memory:
                row_cells = table.add_row().cells
                row_cells[0].text = str(_('Memory')) + ': '
                row_cells[1].text = protocol.device.memory
            # cells width
            for cell in table.columns[0].cells:
                cell.width = Cm(4)
            for cell in table.columns[1].cells:
                cell.width = Cm(14.5)
            p = document.add_paragraph(str(_('Date of testing')) + ': ' +
                                       str(protocol.date_of_start.strftime("%d.%m.%Y")), style='Heading 2')
            if protocol.date_of_finish:
                p.add_run(' - ' + str(protocol.date_of_finish.strftime("%d.%m.%Y")))
            document.add_paragraph(str(_('Photo')), style='Heading 1')
            table = document.add_table(rows=0, cols=1)
            table.style = 'TableGrid'
            row_cells = table.add_row().cells
            row_cells[0].text = str('')

            p = document.add_paragraph()
            run = p.add_run()
            run.add_break(WD_BREAK.PAGE)

        if performance:
            document.add_paragraph(str(_('Performance results')), style='Heading 1')

            p = document.add_paragraph()
            run = p.add_run()
            run.add_break(WD_BREAK.PAGE)

        if results_table:
            results = protocol.get_results(headers=True)
            document.add_paragraph(str(_('Test results')), style='Heading 1')
            document.add_paragraph(str(_('Testing was carried out in accordance with testplan of ')) +
                                   protocol.testplan.name + ' (' + str(_('document version')) + ': ' +
                                   protocol.testplan.version + ').', style='Normal')
            p = document.add_paragraph(str(_('The table below shows the test results. Description of "Res." column: '
                                             )), style='Normal')
            run = p.add_run(u'\u2713')
            run.bold = True
            run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
            p.add_run(' - ' + str(_('Successful')) + ', ')
            run = p.add_run(u'\u274C')
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            p.add_run(' - ' + str(_('Failed')) + ', ')
            run = p.add_run('?')
            run.bold = True
            p.add_run(' - ' + str(_('Skipped')) + ', ')
            run = p.add_run(u'\u00b1')
            run.bold = True
            p.add_run(' - ' + str(_('Passed with warning')) + '.')

            table = document.add_table(rows=1, cols=4)
            table.style = 'TableGrid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '№'
            hdr_cells[1].text = str(_('Test name'))
            hdr_cells[2].text = str(_('Res.'))
            hdr_cells[3].text = str(_('Comment'))
            for hdr_cell in hdr_cells:
                hdr_cell.paragraphs[0].runs[0].font.bold = True

            for result in results:
                row_cells = table.add_row().cells
                if result['header']:
                    row_cells[1].merge(row_cells[2]).merge(row_cells[3])
                    row_cells[0].text = str(result['num'])
                    row_cells[0].paragraphs[0].runs[0].font.bold = True
                    row_cells[1].text = str(result['category_name'])
                    row_cells[1].paragraphs[0].runs[0].font.bold = True
                else:
                    row_cells[0].text = str(result['num'][0]) + '.' + str(result['num'][1])
                    row_cells[1].text = str(result['test_name'])
                    paragraph = row_cells[2].paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = paragraph.add_run()
                    if result['result'] == 0 or result['result'] is None:
                        run.add_text('?')
                        run.font.bold = True
                    elif result['result'] == 1:
                        run.add_text(u'\u274C')
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                    elif result['result'] == 2:
                        run.add_text(u'\u00b1')
                        run.font.bold = True
                    elif result['result'] == 3:
                        run.add_text(u'\u2713')
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
                    row_cells[3].text = str(result['comment'])
            # cells width
            for cell in table.columns[0].cells:
                cell.width = Cm(1.4)
            for cell in table.columns[1].cells:
                cell.width = Cm(12.5)
            for cell in table.columns[2].cells:
                cell.width = Cm(1)
            for cell in table.columns[3].cells:
                cell.width = Cm(4)
            # page break
            p = document.add_paragraph()
            run = p.add_run()
            run.add_break(WD_BREAK.PAGE)

        if summary:
            issues = protocol.get_issues()
            document.add_paragraph(str(_('Protocol summary')), style='Heading 1')
            p = document.add_paragraph(str(_('During testing of device ')), style='Normal')
            run = p.add_run(protocol.device.model)
            run.bold = True
            p.add_run(str(_(' from manufacturer ')))
            run = p.add_run(protocol.device.vendor.name)
            run.bold = True
            p.add_run(str(_(' with firmware version ')))
            run = p.add_run(protocol.sw)
            run.bold = True
            p.add_run(str(_(' and hardware version ')))
            run = p.add_run(protocol.device.hw)
            run.bold = True
            if issues:
                p.add_run(str(_(', the following issues were found:')))
                for issue in issues:
                    document.add_paragraph(issue['text'], style='List Number')
            else:
                p.add_run(str(_(', issues not found.')))

        if team:
            document.add_paragraph(str(_('Worked on testing')), style='Heading 1')

        ###
        file = os.path.join(settings.MEDIA_ROOT + '/docx_builder/protocols/', 'Protocol_' + str(protocol.id) + '.docx')
        document.save(file)
        if os.path.exists(file):
            with open(file, 'rb') as fh:
                response = HttpResponse(fh.read(), content_type="application/vnd.ms-word")
                response['Content-Disposition'] = 'inline; filename=' + os.path.basename(file)
                return response
        raise Http404

    else:
        message = [False, _('Page not found')]
        return render(request, 'device/message.html', {'message': message})


@login_required
def build_protocol_detailed(request):
    if request.method == 'POST':
        protocol = get_object_or_404(Protocol, id=request.POST['protocol_id'])
        docx_profile = get_object_or_404(DocxProfile, id=request.POST['docx_profile_id'])
        header = test_purpose = test_procedure = test_expected = test_images = test_configs = test_links = False
        result_configs = result_images = result_notes = result_status = False
        try:
            if request.POST['header']:
                header = True
        except MultiValueDictKeyError:
            pass
        try:
            if request.POST['test_purpose']:
                test_purpose = True
        except MultiValueDictKeyError:
            pass
        try:
            if request.POST['test_procedure']:
                test_procedure = True
        except MultiValueDictKeyError:
            pass
        try:
            if request.POST['test_expected']:
                test_expected = True
        except MultiValueDictKeyError:
            pass
        try:
            if request.POST['test_images']:
                test_images = True
        except MultiValueDictKeyError:
            pass
        try:
            if request.POST['test_configs']:
                test_configs = True
        except MultiValueDictKeyError:
            pass
        try:
            if request.POST['test_links']:
                test_links = True
        except MultiValueDictKeyError:
            pass
        try:
            if request.POST['result_configs']:
                result_configs = True
        except MultiValueDictKeyError:
            pass
        try:
            if request.POST['result_images']:
                result_images = True
        except MultiValueDictKeyError:
            pass
        try:
            if request.POST['result_notes']:
                result_notes = True
        except MultiValueDictKeyError:
            pass
        try:
            if request.POST['result_status']:
                result_status = True
        except MultiValueDictKeyError:
            pass

        document = build_document(docx_profile=docx_profile)

        # margin
        section = document.sections
        section[0].left_margin = Cm(2)
        section[0].right_margin = Cm(1)
        section[0].top_margin = Cm(5)
        section[0].bottom_margin = Cm(1)
        ###
        if header:
            style = document.styles.add_style('Header Table', WD_STYLE_TYPE.TABLE)
            style.base_style = document.styles['Table Grid']

            header = document.sections[0].header
            table = header.add_table(rows=0, cols=3, width=Cm(19.5))
            table.style = 'Header Table'
            row_cells = table.add_row().cells
            paragraph = row_cells[0].paragraphs[0]
            run = paragraph.add_run()
            if docx_profile.header_logo:
                run.add_picture(docx_profile.header_logo, width=Inches(1.25))
            paragraph = row_cells[1].paragraphs[0]
            run = paragraph.add_run()
            run.add_text(str(_('Protocol')) + ' ' + protocol.device.vendor.name + ' ' + protocol.device.model)
            if protocol.device.hw:
                run.add_text(' (' + protocol.device.hw + ')')
            row_cells = table.add_row().cells
            paragraph = row_cells[0].paragraphs[0]
            run = paragraph.add_run()
            if docx_profile.header_text1:
                run.add_text(docx_profile.header_text1)
            paragraph = row_cells[1].paragraphs[0]
            run = paragraph.add_run()
            if docx_profile.header_text2:
                run.add_text(docx_profile.header_text2)

            table.cell(0, 1).merge(table.cell(0, 2))
            table.cell(0, 0).width = Cm(5)
            table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.cell(1, 0).width = Cm(5)
            table.cell(1, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.cell(0, 1).width = Cm(10.5)
            table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.cell(1, 1).width = Cm(10.5)
            table.cell(1, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.cell(0, 2).width = Cm(4)
            table.cell(0, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.cell(1, 2).width = Cm(4)
            table.cell(1, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.cell(0, 1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        document.add_paragraph(str(_('Detailed test results')), style='Title')
        results = protocol.get_results(headers=True)
        for result in results:
            if result['header']:
                document.add_heading(str(result['num']) + '. ' + result['category_name'], level=1)
            else:
                document.add_heading(str(result['num'][0]) + '.' + str(result['num'][1]) + '. ' +
                                     result['test_name'], level=2)
                table = document.add_table(rows=0, cols=2)
                table.style = 'TableGrid'
                # test table
                if test_purpose:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(_('Purpose')) + ': '
                    if result['test_purpose']:
                        row_cells[1].text = result['test_purpose']
                if test_procedure:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(_('Procedure')) + ': '
                    if result['test_procedure']:
                        procedure_text = result['test_procedure'].replace('\r', '')
                        row_cells[1].text = procedure_text
                if test_expected:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(_('Expected Result')) + ': '
                    if result['test_expected']:
                        expected_text = result['test_expected'].replace('\r', '')
                        row_cells[1].text = expected_text

                for cell in table.columns[0].cells:
                    cell.width = Cm(4)
                for cell in table.columns[1].cells:
                    cell.width = Cm(14.5)

                if result_configs:
                    if result['configs']:
                        document.add_heading(str(_('Configurations')), level=3)
                        for config in result['configs']:
                            if config['desc']:
                                document.add_paragraph(config['desc'], style='Caption')
                            config_text = config['config'].replace('\r', '')
                            table = document.add_table(rows=1, cols=1)
                            table.style = 'Table Grid'
                            shade_cells([table.cell(0, 0)], "#e3e8ec")
                            table.cell(0, 0).text = config_text
                            table.cell(0, 0).paragraphs[0].style = 'Quote'

                if result_notes:
                    if result['notes']:
                        document.add_heading(str(_('Notes')), level=3)
                        for note in result['notes']:
                            if note['desc']:
                                document.add_paragraph(note['desc'], style='Caption')
                            note_text = note['text'].replace('\r', '')
                            table = document.add_table(rows=1, cols=1)
                            table.style = 'Table Grid'
                            shade_cells([table.cell(0, 0)], "#e3e8ec")
                            table.cell(0, 0).text = note_text
                            table.cell(0, 0).paragraphs[0].style = 'Quote'

        ###
        file = os.path.join(settings.MEDIA_ROOT + '/docx_builder/protocols_detailed/',
                            'Detailed_Protocol_' + str(protocol.id) + '.docx')
        document.save(file)
        if os.path.exists(file):
            with open(file, 'rb') as fh:
                response = HttpResponse(fh.read(), content_type="application/vnd.ms-word")
                response['Content-Disposition'] = 'inline; filename=' + os.path.basename(file)
                return response
        raise Http404
    else:
        message = [False, _('Page not found')]
        return render(request, 'device/message.html', {'message': message})


@login_required
def build_testplan(request):
    if request.method == 'POST':
        testplan = get_object_or_404(TestPlan, id=request.POST['testplan_id'])
        docx_profile = get_object_or_404(DocxProfile, id=request.POST['docx_profile_id'])
        title_page = header = purpose = procedure = expected = configs = images = False
        try:
            if request.POST['title_page']:
                title_page = True
        except MultiValueDictKeyError:
            pass
        try:
            if request.POST['header']:
                header = True
        except MultiValueDictKeyError:
            pass
        try:
            if request.POST['purpose']:
                purpose = True
        except MultiValueDictKeyError:
            pass
        try:
            if request.POST['procedure']:
                procedure = True
        except MultiValueDictKeyError:
            pass
        try:
            if request.POST['expected']:
                expected = True
        except MultiValueDictKeyError:
            pass
        try:
            if request.POST['configs']:
                configs = True
        except MultiValueDictKeyError:
            pass
        try:
            if request.POST['images']:
                images = True
        except MultiValueDictKeyError:
            pass

        document = build_document(docx_profile=docx_profile)
        # margin
        section = document.sections
        section[0].left_margin = Cm(2)
        section[0].right_margin = Cm(1)
        section[0].top_margin = Cm(5)
        section[0].bottom_margin = Cm(1)
        ###
        if header:
            style = document.styles.add_style('Header Table', WD_STYLE_TYPE.TABLE)
            style.base_style = document.styles['Table Grid']

            header = document.sections[0].header
            table = header.add_table(rows=0, cols=3, width=Cm(19.5))
            table.style = 'Header Table'
            row_cells = table.add_row().cells
            paragraph = row_cells[0].paragraphs[0]
            run = paragraph.add_run()
            if docx_profile.header_logo:
                run.add_picture(docx_profile.header_logo, width=Inches(1.25))
            paragraph = row_cells[1].paragraphs[0]
            run = paragraph.add_run()
            run.add_text(str(_('Testplan')) + ' ' + testplan.name)
            row_cells = table.add_row().cells
            paragraph = row_cells[0].paragraphs[0]
            run = paragraph.add_run()
            run.add_text(str(_('document version')) + ': ' + testplan.version)
            paragraph = row_cells[1].paragraphs[0]
            run = paragraph.add_run()
            if docx_profile.header_text2:
                run.add_text(docx_profile.header_text2)

            table.cell(0, 1).merge(table.cell(0, 2))
            table.cell(0, 0).width = Cm(5)
            table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.cell(1, 0).width = Cm(5)
            table.cell(1, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.cell(0, 1).width = Cm(10.5)
            table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.cell(1, 1).width = Cm(10.5)
            table.cell(1, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.cell(0, 2).width = Cm(4)
            table.cell(0, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.cell(1, 2).width = Cm(4)
            table.cell(1, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.cell(0, 1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        ###
        document.add_paragraph(str(_('Testplan')), style='Title')

        categories = Category.objects.filter(testplan=testplan).order_by('priority')
        for i, category in enumerate(categories):
            document.add_heading(str(i+1) + '. ' + category.name, level=1)
            tests = Test.objects.filter(cat=category).order_by('priority')
            for j, test in enumerate(tests):
                document.add_heading(str(i + 1) + '.' + str(j + 1) + '. ' + test.name, level=2)
                # test table
                table = document.add_table(rows=0, cols=2)
                table.style = 'TableGrid'

                if purpose:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(_('Purpose')) + ': '
                    if test.purpose:
                        row_cells[1].text = test.purpose
                if procedure:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(_('Procedure')) + ': '
                    if test.procedure:
                        row_cells[1].text = textile_to_docx(test.procedure)

                if expected:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(_('Expected result')) + ': '
                    if test.expected:
                        row_cells[1].text = textile_to_docx(test.expected)

                for cell in table.columns[0].cells:
                    cell.width = Cm(4)
                for cell in table.columns[1].cells:
                    cell.width = Cm(15)

                if configs:
                    test_configs = TestConfig.objects.filter(test=test).order_by('id')
                    if test_configs:
                        document.add_heading(str(_('Configurations')), level=3)
                        for test_config in test_configs:
                            if test_config.parent:
                                document.add_paragraph(test_config.parent.desc, style='Caption')
                            else:
                                document.add_paragraph(test_config.desc, style='Caption')
                            table = document.add_table(rows=1, cols=1)
                            table.style = 'Table Grid'
                            shade_cells([table.cell(0, 0)], "#e3e8ec")
                            if test_config.parent:
                                test_config.parent.config = test_config.parent.config.replace('\r', '')
                                table.cell(0, 0).text = test_config.parent.config
                            else:
                                test_config.config = test_config.config.replace('\r', '')
                                table.cell(0, 0).text = test_config.config
                if images:
                    for test_image in TestImage.objects.filter(test=test).order_by('id'):
                        if test_image.parent:
                            if test_image.parent.desc:
                                document.add_heading(test_image.parent.desc, level=3)
                            if test_image.parent.image:
                                document.add_picture(test_image.parent.image, width=Cm(17))
                        else:
                            if test_image.desc:
                                document.add_heading(test_image.desc, level=3)
                            if test_image.image:
                                document.add_picture(test_image.image, width=Cm(17))

        ###
        file = os.path.join(settings.MEDIA_ROOT + '/docx_builder/testplans/',
                            'Testplan_' + str(testplan.id) + '.docx')
        document.save(file)
        if os.path.exists(file):
            with open(file, 'rb') as fh:
                response = HttpResponse(fh.read(), content_type="application/vnd.ms-word")
                response['Content-Disposition'] = 'inline; filename=' + os.path.basename(file)
                return response
        raise Http404
    else:
        message = [False, _('Page not found')]
        return render(request, 'device/message.html', {'message': message})


def build_document(docx_profile: DocxProfile):
    document = Document()
    document.styles['Title'].font.name = docx_profile.title_font_name
    document.styles['Title'].font.color.rgb = RGBColor(docx_profile.title_font_color_red,
                                                       docx_profile.title_font_color_green,
                                                       docx_profile.title_font_color_blue)
    document.styles['Title'].font.size = Pt(docx_profile.title_font_size)
    document.styles['Title'].font.bold = docx_profile.title_font_bold
    document.styles['Title'].font.italic = docx_profile.title_font_italic
    document.styles['Title'].font.underline = docx_profile.title_font_underline
    document.styles['Title'].paragraph_format.space_before = Pt(docx_profile.title_space_before)
    document.styles['Title'].paragraph_format.space_after = Pt(docx_profile.title_space_after)
    if docx_profile.title_alignment == 0:
        document.styles['Title'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    elif docx_profile.title_alignment == 1:
        document.styles['Title'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif docx_profile.title_alignment == 2:
        document.styles['Title'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    elif docx_profile.title_alignment == 3:
        document.styles['Title'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    document.styles['Heading 1'].font.name = docx_profile.h1_font_name
    document.styles['Heading 1'].font.color.rgb = RGBColor(docx_profile.h1_font_color_red,
                                                           docx_profile.h1_font_color_green,
                                                           docx_profile.h1_font_color_blue)
    document.styles['Heading 1'].font.size = Pt(docx_profile.h1_font_size)
    document.styles['Heading 1'].font.bold = docx_profile.h1_font_bold
    document.styles['Heading 1'].font.italic = docx_profile.h1_font_italic
    document.styles['Heading 1'].font.underline = docx_profile.h1_font_underline
    document.styles['Heading 1'].paragraph_format.space_before = Pt(docx_profile.h1_space_before)
    document.styles['Heading 1'].paragraph_format.space_after = Pt(docx_profile.h1_space_after)
    if docx_profile.h1_alignment == 0:
        document.styles['Heading 1'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    elif docx_profile.h1_alignment == 1:
        document.styles['Heading 1'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif docx_profile.h1_alignment == 2:
        document.styles['Heading 1'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    elif docx_profile.h1_alignment == 3:
        document.styles['Heading 1'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    document.styles['Heading 2'].font.name = docx_profile.h2_font_name
    document.styles['Heading 2'].font.color.rgb = RGBColor(docx_profile.h2_font_color_red,
                                                           docx_profile.h2_font_color_green,
                                                           docx_profile.h2_font_color_blue)
    document.styles['Heading 2'].font.size = Pt(docx_profile.h2_font_size)
    document.styles['Heading 2'].font.bold = docx_profile.h2_font_bold
    document.styles['Heading 2'].font.italic = docx_profile.h2_font_italic
    document.styles['Heading 2'].font.underline = docx_profile.h2_font_underline
    document.styles['Heading 2'].paragraph_format.space_before = Pt(docx_profile.h2_space_before)
    document.styles['Heading 2'].paragraph_format.space_after = Pt(docx_profile.h2_space_after)
    if docx_profile.h2_alignment == 0:
        document.styles['Heading 2'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    elif docx_profile.h2_alignment == 1:
        document.styles['Heading 2'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif docx_profile.h2_alignment == 2:
        document.styles['Heading 2'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    elif docx_profile.h2_alignment == 3:
        document.styles['Heading 2'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    document.styles['Heading 3'].font.name = docx_profile.h3_font_name
    document.styles['Heading 3'].font.color.rgb = RGBColor(docx_profile.h3_font_color_red,
                                                           docx_profile.h3_font_color_green,
                                                           docx_profile.h3_font_color_blue)
    document.styles['Heading 3'].font.size = Pt(docx_profile.h3_font_size)
    document.styles['Heading 3'].font.bold = docx_profile.h3_font_bold
    document.styles['Heading 3'].font.italic = docx_profile.h3_font_italic
    document.styles['Heading 3'].font.underline = docx_profile.h3_font_underline
    document.styles['Heading 3'].paragraph_format.space_before = Pt(docx_profile.h3_space_before)
    document.styles['Heading 3'].paragraph_format.space_after = Pt(docx_profile.h3_space_after)
    if docx_profile.h3_alignment == 0:
        document.styles['Heading 3'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    elif docx_profile.h3_alignment == 1:
        document.styles['Heading 3'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif docx_profile.h3_alignment == 2:
        document.styles['Heading 3'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    elif docx_profile.h3_alignment == 3:
        document.styles['Heading 3'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    document.styles['Normal'].font.name = docx_profile.normal_font_name
    document.styles['Normal'].font.color.rgb = RGBColor(docx_profile.normal_font_color_red,
                                                        docx_profile.normal_font_color_green,
                                                        docx_profile.normal_font_color_blue)
    document.styles['Normal'].font.size = Pt(docx_profile.normal_font_size)
    document.styles['Normal'].font.bold = docx_profile.normal_font_bold
    document.styles['Normal'].font.italic = docx_profile.normal_font_italic
    document.styles['Normal'].font.underline = docx_profile.normal_font_underline
    document.styles['Normal'].paragraph_format.space_before = Pt(docx_profile.normal_space_before)
    document.styles['Normal'].paragraph_format.space_after = Pt(docx_profile.normal_space_after)
    if docx_profile.normal_alignment == 0:
        document.styles['Normal'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    elif docx_profile.normal_alignment == 1:
        document.styles['Normal'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif docx_profile.normal_alignment == 2:
        document.styles['Normal'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    elif docx_profile.normal_alignment == 3:
        document.styles['Normal'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    document.styles['Caption'].font.name = docx_profile.caption_font_name
    document.styles['Caption'].font.color.rgb = RGBColor(docx_profile.caption_font_color_red,
                                                         docx_profile.caption_font_color_green,
                                                         docx_profile.caption_font_color_blue)
    document.styles['Caption'].font.size = Pt(docx_profile.caption_font_size)
    document.styles['Caption'].font.bold = docx_profile.caption_font_bold
    document.styles['Caption'].font.italic = docx_profile.caption_font_italic
    document.styles['Caption'].font.underline = docx_profile.caption_font_underline
    document.styles['Caption'].paragraph_format.space_before = Pt(docx_profile.caption_space_before)
    document.styles['Caption'].paragraph_format.space_after = Pt(docx_profile.caption_space_after)
    if docx_profile.caption_alignment == 0:
        document.styles['Caption'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    elif docx_profile.caption_alignment == 1:
        document.styles['Caption'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif docx_profile.caption_alignment == 2:
        document.styles['Caption'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    elif docx_profile.caption_alignment == 3:
        document.styles['Caption'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    document.styles['Quote'].font.name = docx_profile.quote_font_name
    document.styles['Quote'].font.color.rgb = RGBColor(docx_profile.quote_font_color_red,
                                                       docx_profile.quote_font_color_green,
                                                       docx_profile.quote_font_color_blue)
    document.styles['Quote'].font.size = Pt(docx_profile.quote_font_size)
    document.styles['Quote'].font.bold = docx_profile.quote_font_bold
    document.styles['Quote'].font.italic = docx_profile.quote_font_italic
    document.styles['Quote'].font.underline = docx_profile.quote_font_underline
    document.styles['Quote'].paragraph_format.space_before = Pt(docx_profile.quote_space_before)
    document.styles['Quote'].paragraph_format.space_after = Pt(docx_profile.quote_space_after)
    if docx_profile.quote_alignment == 0:
        document.styles['Quote'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    elif docx_profile.quote_alignment == 1:
        document.styles['Quote'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif docx_profile.quote_alignment == 2:
        document.styles['Quote'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    elif docx_profile.quote_alignment == 3:
        document.styles['Quote'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    return document


def shade_cells(cells, shade):
    for cell in cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcVAlign = OxmlElement("w:shd")
        tcVAlign.set(qn("w:fill"), shade)
        tcPr.append(tcVAlign)


def textile_to_docx(foo: str):
    j = 1
    k = 1
    lines = foo.split('\r')
    for i, line in enumerate(lines):
        if lines[i].startswith('# '):
            lines[i] = str(j) + '.' + lines[i][1:]
            j += 1
        elif lines[i].startswith('\n# '):
            lines[i] = '\n' + str(j) + '.' + lines[i][2:]
            j += 1
            k = 1
        elif lines[i].startswith('\n## '):
            lines[i] = '\n    ' + str(j-1) + '.' + str(k) + '.' + lines[i][3:]
            k += 1
        elif lines[i].startswith('\n* '):
            lines[i] = '\n' + u'\u2022' + ' ' + lines[i][2:]
        elif lines[i].startswith('\n** '):
            lines[i] = '\n    ' + u'\u2022' + ' ' + lines[i][3:]
    bar = ''.join(lines)
    return bar
